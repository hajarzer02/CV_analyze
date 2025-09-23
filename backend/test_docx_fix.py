#!/usr/bin/env python3
"""
Test the DOCX extraction fix with various scenarios.
"""

import os
import sys
from docx import Document

def test_docx_fix():
    """Test the DOCX extraction fix."""
    print("🔍 Testing DOCX Extraction Fix")
    print("=" * 50)
    
    # Test 1: Normal DOCX file
    print("\n1. Testing normal DOCX file:")
    test_content = """
John Smith
Software Engineer
Email: john.smith@example.com
Phone: +1-555-123-4567

PROFESSIONAL SUMMARY
Experienced software engineer with 5+ years of experience.

SKILLS
- Python
- JavaScript
- React
- Node.js

EXPERIENCE
Senior Software Engineer at TechCorp (2020-2024)
- Led development of web applications
- Mentored junior developers

EDUCATION
Bachelor of Computer Science
University of Technology (2016-2020)
"""
    
    doc = Document()
    for line in test_content.strip().split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    
    normal_docx = "test_normal.docx"
    doc.save(normal_docx)
    
    try:
        from cv_extractor_cli import load_text
        raw_text = load_text(normal_docx)
        print(f"   Raw text length: {len(raw_text)} characters")
        print(f"   First 200 characters: '{raw_text[:200]}'")
        
        if len(raw_text) > 0:
            print("   ✅ Normal DOCX extraction successful")
        else:
            print("   ❌ Normal DOCX extraction failed")
    except Exception as e:
        print(f"   ❌ Error: {e}")
    finally:
        if os.path.exists(normal_docx):
            os.remove(normal_docx)
    
    # Test 2: Empty DOCX file
    print("\n2. Testing empty DOCX file:")
    empty_doc = Document()
    empty_docx = "test_empty.docx"
    empty_doc.save(empty_docx)
    
    try:
        from cv_extractor_cli import load_text
        raw_text = load_text(empty_docx)
        print(f"   Raw text length: {len(raw_text)} characters")
        
        if len(raw_text) == 0:
            print("   ✅ Empty DOCX handled correctly")
        else:
            print("   ⚠️  Empty DOCX returned content (unexpected)")
    except Exception as e:
        print(f"   ❌ Error: {e}")
    finally:
        if os.path.exists(empty_docx):
            os.remove(empty_docx)
    
    # Test 3: DOCX with only tables
    print("\n3. Testing DOCX with tables:")
    table_doc = Document()
    table = table_doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "John Doe"
    table.cell(1, 0).text = "Email"
    table.cell(1, 1).text = "john.doe@example.com"
    table.cell(2, 0).text = "Phone"
    table.cell(2, 1).text = "+1-555-123-4567"
    
    table_docx = "test_table.docx"
    table_doc.save(table_docx)
    
    try:
        from cv_extractor_cli import load_text
        raw_text = load_text(table_docx)
        print(f"   Raw text length: {len(raw_text)} characters")
        print(f"   Content: '{raw_text}'")
        
        if "John Doe" in raw_text:
            print("   ✅ Table content extracted successfully")
        else:
            print("   ❌ Table content not extracted")
    except Exception as e:
        print(f"   ❌ Error: {e}")
    finally:
        if os.path.exists(table_docx):
            os.remove(table_docx)
    
    # Test 4: Test LLaMA service with empty text
    print("\n4. Testing LLaMA service with empty text:")
    try:
        sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'ai-service'))
        from llama_service import LlamaService
        
        llama_service = LlamaService()
        result = llama_service.structure_cv_text("")
        
        if "extraction_error" in result:
            print("   ✅ LLaMA service correctly handled empty text")
            print(f"   Error message: {result.get('extraction_error')}")
        else:
            print("   ⚠️  LLaMA service did not handle empty text properly")
    except Exception as e:
        print(f"   ❌ Error: {e}")
    
    print("\n" + "=" * 50)
    print("✅ DOCX extraction fix testing completed!")

if __name__ == "__main__":
    test_docx_fix()


