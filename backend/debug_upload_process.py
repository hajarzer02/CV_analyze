#!/usr/bin/env python3
"""
Debug the exact upload process to find why DOCX extraction fails.
"""

import os
import sys
import shutil
from docx import Document

def debug_upload_process():
    """Debug the upload process step by step."""
    print("🔍 Debugging Upload Process")
    print("=" * 50)
    
    # Create a test DOCX file
    test_content = """
Doha Bouhali
Développeuse Full Stack
Email: doha.bouhali@example.com
Téléphone: +212 6 50 11 31 97
Adresse: Mers Sultan Casablanca 20250, Maroc

PROFIL PROFESSIONNEL
Développeuse passionnée avec une expertise en développement web et mobile.

COMPÉTENCES TECHNIQUES
- Langages: Python, JavaScript, PHP, Dart
- Frameworks: React, Laravel, Flutter
- Bases de données: MySQL, Oracle
- Outils: Git, Docker, Linux

EXPÉRIENCE
Développeuse Full Stack chez TechCorp (2022-2024)
- Développement d'applications web avec React et Node.js
- Création d'APIs REST avec Laravel
- Gestion de bases de données MySQL

Stagiaire Développeuse chez StartupXYZ (2021-2022)
- Développement frontend avec React
- Intégration d'APIs externes

FORMATION
Licence Génie Informatique
IGA Casablanca (2019-2022)
"""
    
    # Create test DOCX file
    doc = Document()
    for line in test_content.strip().split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    
    test_docx_path = "test_upload_cv.docx"
    doc.save(test_docx_path)
    print(f"✅ Created test DOCX file: {test_docx_path}")
    
    try:
        # Simulate the exact upload process from main.py
        print("\n1. Simulating upload process:")
        
        # Step 1: Save uploaded file (like in main.py)
        uploads_dir = "uploads"
        os.makedirs(uploads_dir, exist_ok=True)
        
        uploaded_file_path = os.path.join(uploads_dir, "CV_Doha_ENG.docx")
        print(f"   Copying file to: {uploaded_file_path}")
        
        shutil.copy2(test_docx_path, uploaded_file_path)
        
        # Check if file exists and has content
        if os.path.exists(uploaded_file_path):
            file_size = os.path.getsize(uploaded_file_path)
            print(f"   ✅ File copied successfully, size: {file_size} bytes")
        else:
            print("   ❌ File copy failed")
            return
        
        # Step 2: Extract raw text (like in main.py)
        print("\n2. Extracting raw text:")
        try:
            from cv_extractor_cli import CVExtractor
            cv_extractor = CVExtractor()
            
            raw_text = cv_extractor.extract_raw_text(uploaded_file_path)
            print(f"   Raw text length: {len(raw_text)} characters")
            print(f"   First 300 characters:")
            print(f"   '{raw_text[:300]}'")
            
            if len(raw_text) == 0:
                print("   ❌ ERROR: Raw text is empty!")
                
                # Debug the file
                print("\n   Debugging empty text issue:")
                try:
                    doc = Document(uploaded_file_path)
                    print(f"   Document paragraphs: {len(doc.paragraphs)}")
                    
                    for i, p in enumerate(doc.paragraphs[:5]):
                        print(f"   Paragraph {i}: '{p.text}'")
                        
                except Exception as e:
                    print(f"   Error reading DOCX: {e}")
            else:
                print("   ✅ Raw text extraction successful")
                
        except Exception as e:
            print(f"   ❌ Error in raw text extraction: {e}")
            import traceback
            traceback.print_exc()
        
        # Step 3: Test LLaMA processing
        print("\n3. Testing LLaMA processing:")
        try:
            sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'ai-service'))
            from llama_service import LlamaService
            
            llama_service = LlamaService()
            print(f"   Provider: {llama_service.provider}")
            
            if len(raw_text) > 0:
                extracted_data = llama_service.structure_cv_text(raw_text)
                print(f"   LLaMA result contact info: {extracted_data.get('contact_info', {})}")
                
                # Check content preservation
                preservation_report = llama_service.verify_content_preservation(raw_text, extracted_data)
                print(f"   Content preservation score: {preservation_report['content_preservation_score']:.2f}")
                
                if preservation_report['content_preservation_score'] < 0.1:
                    print("   ⚠️  WARNING: Very low content preservation score!")
                    print("   This explains why John Doe is returned")
                else:
                    print("   ✅ Good content preservation")
            else:
                print("   ❌ Cannot process with LLaMA - no raw text")
                
        except Exception as e:
            print(f"   ❌ Error in LLaMA processing: {e}")
            import traceback
            traceback.print_exc()
        
        # Step 4: Check the data.txt file
        print("\n4. Checking data.txt file:")
        data_txt_path = os.path.join(uploads_dir, "data.txt")
        if os.path.exists(data_txt_path):
            with open(data_txt_path, 'r', encoding='utf-8') as f:
                data_content = f.read()
            print(f"   data.txt length: {len(data_content)} characters")
            if len(data_content) > 0:
                print(f"   First 200 characters: '{data_content[:200]}'")
            else:
                print("   ❌ data.txt is empty!")
        else:
            print("   ❌ data.txt file not found!")
        
    except Exception as e:
        print(f"❌ Error during debugging: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # Clean up
        for file_path in [test_docx_path, uploaded_file_path]:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"   🧹 Cleaned up: {file_path}")

if __name__ == "__main__":
    debug_upload_process()






