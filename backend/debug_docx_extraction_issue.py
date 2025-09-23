#!/usr/bin/env python3
"""
Debug the specific DOCX extraction issue that's causing empty text.
"""

import os
import sys
from docx import Document

def debug_docx_extraction():
    """Debug DOCX extraction issue."""
    print("🔍 Debugging DOCX Extraction Issue")
    print("=" * 50)
    
    # Create a test DOCX file similar to the one that's failing
    test_content = """
Doha Bouhali
Développeuse Full Stack
Email: doha.bouhali@example.com
Téléphone: +212 6 50 11 31 97
Adresse: Mers Sultan Casablanca 20250, Maroc

PROFIL PROFESSIONNEL
Développeuse passionnée avec une expertise en développement web et mobile.

COMPÉTENCES TECHNIQUES
- Langages: Python, JavaScript, PHP, Dart
- Frameworks: React, Laravel, Flutter
- Bases de données: MySQL, Oracle
- Outils: Git, Docker, Linux

EXPÉRIENCE
Développeuse Full Stack chez TechCorp (2022-2024)
- Développement d'applications web avec React et Node.js
- Création d'APIs REST avec Laravel
- Gestion de bases de données MySQL

Stagiaire Développeuse chez StartupXYZ (2021-2022)
- Développement frontend avec React
- Intégration d'APIs externes

FORMATION
Licence Génie Informatique
IGA Casablanca (2019-2022)
"""
    
    # Create test DOCX file
    doc = Document()
    for line in test_content.strip().split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    
    test_docx_path = "debug_test_cv.docx"
    doc.save(test_docx_path)
    print(f"✅ Created test DOCX file: {test_docx_path}")
    
    try:
        # Test 1: Direct DOCX extraction
        print("\n1. Testing direct DOCX extraction:")
        try:
            doc = Document(test_docx_path)
            lines = []
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    lines.append(t)
            
            raw_text = "\n".join(lines)
            print(f"   Raw text length: {len(raw_text)} characters")
            print(f"   First 300 characters:")
            print(f"   '{raw_text[:300]}'")
            
            if len(raw_text) == 0:
                print("   ❌ ERROR: Raw text is empty!")
            else:
                print("   ✅ Raw text extraction successful")
                
        except Exception as e:
            print(f"   ❌ Error in direct extraction: {e}")
            import traceback
            traceback.print_exc()
        
        # Test 2: Using the CV extractor
        print("\n2. Testing CV extractor:")
        try:
            from cv_extractor_cli import load_text, extract_text_from_docx
            
            # Test the specific function
            extracted_text = extract_text_from_docx(test_docx_path)
            print(f"   Extracted text length: {len(extracted_text)} characters")
            print(f"   First 300 characters:")
            print(f"   '{extracted_text[:300]}'")
            
            if len(extracted_text) == 0:
                print("   ❌ ERROR: CV extractor returned empty text!")
            else:
                print("   ✅ CV extractor extraction successful")
                
        except Exception as e:
            print(f"   ❌ Error in CV extractor: {e}")
            import traceback
            traceback.print_exc()
        
        # Test 3: Using load_text function
        print("\n3. Testing load_text function:")
        try:
            loaded_text = load_text(test_docx_path)
            print(f"   Loaded text length: {len(loaded_text)} characters")
            print(f"   First 300 characters:")
            print(f"   '{loaded_text[:300]}'")
            
            if len(loaded_text) == 0:
                print("   ❌ ERROR: load_text returned empty text!")
            else:
                print("   ✅ load_text extraction successful")
                
        except Exception as e:
            print(f"   ❌ Error in load_text: {e}")
            import traceback
            traceback.print_exc()
        
        # Test 4: Check if the file is corrupted
        print("\n4. Checking file integrity:")
        try:
            file_size = os.path.getsize(test_docx_path)
            print(f"   File size: {file_size} bytes")
            
            if file_size == 0:
                print("   ❌ ERROR: File is empty!")
            else:
                print("   ✅ File has content")
                
            # Try to open with Document again
            doc2 = Document(test_docx_path)
            paragraph_count = len(doc2.paragraphs)
            print(f"   Paragraph count: {paragraph_count}")
            
            if paragraph_count == 0:
                print("   ❌ ERROR: No paragraphs found in DOCX!")
            else:
                print("   ✅ Paragraphs found")
                
        except Exception as e:
            print(f"   ❌ Error checking file integrity: {e}")
        
    except Exception as e:
        print(f"❌ Error during debugging: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # Clean up
        if os.path.exists(test_docx_path):
            os.remove(test_docx_path)
            print(f"\n🧹 Cleaned up test file: {test_docx_path}")

if __name__ == "__main__":
    debug_docx_extraction()


